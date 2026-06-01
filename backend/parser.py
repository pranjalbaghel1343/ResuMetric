import pdfplumber
import docx
import os
import re


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file using python-docx."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text.strip()


def extract_text(file_path: str) -> str:
    """Auto-detect format and extract text."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    return ""


def extract_name_from_text(text: str, filename: str) -> str:
    """Try to extract candidate name from resume text."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    # The first non-empty line is often the name
    if lines:
        first_line = lines[0]
        # If it's reasonably short and looks like a name
        if 2 <= len(first_line.split()) <= 5 and len(first_line) < 50:
            return first_line
    # Fallback to filename (without extension)
    name = os.path.splitext(filename)[0]
    name = re.sub(r'[_\-\.]', ' ', name).title()
    return name
